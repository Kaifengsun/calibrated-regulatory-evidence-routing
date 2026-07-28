"""Build the editable Word manuscript from the reviewed Markdown source."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "manuscript" / "manuscript.md"
REFERENCES = ROOT / "manuscript" / "references.md"
OUTPUT = ROOT / "output" / "word" / "When_Does_Evidence_Expansion_Help_Revised.docx"

INK = "18324A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
HEADER_FILL = "F4F6F9"
GRID = "B7C4CF"
MUTED = "5B6570"


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa, strict=True):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_borders(cell, color=GRID, size=4):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_run_font(run, name="Calibri", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (fld_char, instr, sep, text, end):
        run._r.append(node)


def set_image_alt(inline_shape, description):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description)


def strip_markdown(text):
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = text.replace("**", "").replace("`", "")
    return text


def add_inline(paragraph, text, base_size=11):
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=base_size)
        markup_fragment = match.group(0)
        if markup_fragment.startswith("**"):
            run = paragraph.add_run(markup_fragment[2:-2])
            set_run_font(run, size=base_size, bold=True)
        elif markup_fragment.startswith("*"):
            run = paragraph.add_run(markup_fragment[1:-1])
            set_run_font(run, size=base_size, italic=True)
        else:
            run = paragraph.add_run(markup_fragment[1:-1])
            set_run_font(run, name="Consolas", size=base_size - 0.5)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=base_size)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.widow_control = True

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.page_break_before = False

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9.5)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string(INK)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption.paragraph_format.space_before = Pt(8)
    caption.paragraph_format.space_after = Pt(5)
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.keep_with_next = True


def configure_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("RESEARCH ARTICLE  |  REGULATORY EVIDENCE RETRIEVAL")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    add_field(p, "PAGE")


def add_title_block(doc, title, author, affiliation):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(title)
    set_run_font(run, size=19, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(author)
    set_run_font(run, size=11.5, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(affiliation)
    set_run_font(run, size=10.5, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.keep_with_next = True
    run = p.add_run("Sole author and corresponding author")
    set_run_font(run, size=9.5, italic=True, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    p.paragraph_format.keep_with_next = True
    run = p.add_run("Editable pre-submission manuscript")
    set_run_font(run, size=9.5, italic=True, color=MUTED)


def parse_table(lines, start):
    rows = []
    index = start
    while index < len(lines) and lines[index].startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        rows.append(cells)
        index += 1
    if len(rows) >= 2 and all(set(cell) <= {"-", ":"} for cell in rows[1]):
        rows.pop(1)
    return rows, index


def table_widths(rows):
    cols = len(rows[0])
    if cols == 7:
        return [1870, 1300, 1050, 1520, 700, 600, 2320]
    if cols == 5:
        return [2300, 1760, 1420, 1820, 2060]
    if cols == 4:
        return [1600, 2600, 2580, 2580]
    weights = [max(8, max(len(row[i]) for row in rows)) for i in range(cols)]
    total = sum(weights)
    widths = [int(9360 * weight / total) for weight in weights]
    widths[-1] += 9360 - sum(widths)
    return widths


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    widths = table_widths(rows)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for r_index, (row, values) in enumerate(zip(table.rows, rows, strict=True)):
        for c_index, (cell, value) in enumerate(zip(row.cells, values, strict=True)):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_borders(cell)
            if r_index == 0:
                set_cell_fill(cell, HEADER_FILL)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
                if c_index == 0
                else WD_ALIGN_PARAGRAPH.CENTER
            )
            run = p.add_run(strip_markdown(value))
            set_run_font(run, size=8.2 if len(rows[0]) >= 7 else 8.8, bold=(r_index == 0))
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_figure(doc, path, alt):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_together = True
    shape = p.add_run().add_picture(str(path), width=Inches(6.25))
    set_image_alt(shape, alt)


def collect_references():
    lines = REFERENCES.read_text(encoding="utf-8").splitlines()
    refs = []
    current = []
    for line in lines:
        if re.match(r"^\d+\.\s", line):
            if current:
                refs.append(" ".join(current))
            current = [re.sub(r"^\d+\.\s*", "", line)]
        elif current and line.strip():
            current.append(line.strip())
    if current:
        refs.append(" ".join(current))
    return refs


def add_references(doc):
    for number, ref in enumerate(collect_references(), start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.28)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(f"[{number}] ")
        set_run_font(run, size=9.5, bold=True)
        add_inline(p, ref, base_size=9.5)


def build():
    doc = Document()
    configure_styles(doc)
    configure_header_footer(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    title = lines[0].lstrip("# ").strip()
    author = lines[2].strip()
    affiliation = lines[3].strip()
    add_title_block(doc, title, author, affiliation)
    lines = lines[5:]

    paragraph_buffer = []
    skip_reference_placeholder = False

    def flush():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            text = " ".join(part.strip() for part in paragraph_buffer)
            p = doc.add_paragraph()
            add_inline(p, text)
            paragraph_buffer = []

    i = 0
    while i < len(lines):
        line = lines[i]
        if skip_reference_placeholder and line.strip():
            i += 1
            continue
        if not line.strip():
            flush()
            i += 1
            continue
        if line.startswith("## "):
            flush()
            heading = line[3:].strip()
            doc.add_heading(heading, level=1)
            skip_reference_placeholder = heading == "References"
            if skip_reference_placeholder:
                add_references(doc)
            i += 1
            continue
        if line.startswith("### "):
            flush()
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        if line.startswith("**Table ") or line.startswith("**Figure "):
            flush()
            caption_lines = [line]
            while not caption_lines[-1].endswith("**") and i + 1 < len(lines):
                i += 1
                caption_lines.append(lines[i])
            caption_text = strip_markdown(" ".join(caption_lines))
            p = doc.add_paragraph(style="Caption")
            run = p.add_run(caption_text)
            set_run_font(run, size=9.5, bold=True, color=INK)
            i += 1
            continue
        image_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line)
        if image_match:
            flush()
            image_path = (SOURCE.parent / image_match.group(2)).resolve()
            add_figure(doc, image_path, image_match.group(1))
            i += 1
            continue
        if line.startswith("|"):
            flush()
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        paragraph_buffer.append(line)
        i += 1
    flush()

    core = doc.core_properties
    core.title = title
    core.subject = "Cross-domain regulatory evidence retrieval"
    core.author = "Kaifeng Sun"
    core.keywords = (
        "regulatory retrieval; risk-sensitive evaluation; structural context "
        "expansion; selective routing; evidence sufficiency; retrieval-augmented generation"
    )
    core.comments = "Generated from the prespecified Pilot artifacts."
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
